import argparse
import os
import re
import itertools
from collections import defaultdict
from concurrent.futures import ProcessPoolExecutor
import yaml
import pandas as pd
from tqdm import tqdm
from Bio.Align import PairwiseAligner

from stock_phrases import classify_match_row

_WORKER_STATE = {}

# ---------------------------------------------------------------------------
# 1. TEXT PREPROCESSING
# ---------------------------------------------------------------------------

MARKER_PATTERN = re.compile(r'\[\\\[[^\]]*\\\]\]')
METADATA_PATTERN = re.compile(r'@##|@#|/_|/|\{[^}]*\}')
SYLLABLE_SEPARATORS = frozenset(' \t/\n')

# Private Use Area used to encode syllable tokens for BioPython alignment.
_PUA_START = 0xE000
_PUA_END = 0xF8FF
_PUA_SIZE = _PUA_END - _PUA_START + 1


def _syllable_index_to_chars(index: int) -> str:
    """Encode a syllable index as two fixed-width PUA characters."""
    if index >= _PUA_SIZE * _PUA_SIZE:
        raise ValueError(
            f"Too many unique syllables ({index + 1}) for two-character encoding"
        )
    hi = index // _PUA_SIZE
    lo = index % _PUA_SIZE
    return chr(_PUA_START + hi) + chr(_PUA_START + lo)


def tokenize_ewts_syllables(text, strip_newlines=True, strip_markers=True, strip_chars=""):
    """
    Split EWTS text into syllable tokens with original-text spans.

    Syllables are delimited by whitespace and slashes (EWTS convention).
    """
    masked = [False] * len(text)

    if strip_markers:
        for m in MARKER_PATTERN.finditer(text):
            for i in range(m.start(), m.end()):
                masked[i] = True
        for m in METADATA_PATTERN.finditer(text):
            for i in range(m.start(), m.end()):
                masked[i] = True

    syllables = []
    i = 0
    n = len(text)
    while i < n:
        if masked[i]:
            i += 1
            continue
        if strip_newlines and text[i] == '\n':
            i += 1
            continue
        if text[i] in SYLLABLE_SEPARATORS:
            i += 1
            continue

        start = i
        while i < n and not masked[i]:
            if strip_newlines and text[i] == '\n':
                break
            if text[i] in SYLLABLE_SEPARATORS:
                break
            i += 1

        syllable = text[start:i]
        if strip_chars:
            syllable = "".join(c for c in syllable if c not in strip_chars)
        if syllable:
            syllables.append((syllable, start, i))

    return syllables


def build_syllable_encoding(*token_lists):
    """
    Map unique syllables across one or more token lists to fixed-width PUA pairs.
    """
    vocab = sorted({token for slist in token_lists for token in slist})
    if len(vocab) >= _PUA_SIZE * _PUA_SIZE:
        raise ValueError(
            f"Too many unique syllables ({len(vocab)}) for two-character encoding"
        )
    return {token: _syllable_index_to_chars(i) for i, token in enumerate(vocab)}


def build_alignment_map(text, strip_spaces=True, strip_newlines=True, strip_markers=True, strip_chars=""):
    """
    Build a dense string suitable for alignment and a coordinate map back to the original text.
    """
    masked = [False] * len(text)

    if strip_markers:
        for m in MARKER_PATTERN.finditer(text):
            for i in range(m.start(), m.end()):
                masked[i] = True
        for m in METADATA_PATTERN.finditer(text):
            for i in range(m.start(), m.end()):
                masked[i] = True

    coord_map = []
    chars = []
    for i, ch in enumerate(text):
        if masked[i]:
            continue
        if strip_spaces and ch == ' ':
            continue
        if strip_newlines and ch == '\n':
            continue
        if ch in strip_chars:
            continue
        chars.append(ch)
        coord_map.append(i)

    dense_text = "".join(chars)
    return dense_text, coord_map


def build_syllable_alignment_map(text, strip_newlines=True, strip_markers=True, strip_chars=""):
    """
    Build a dense string where each character represents one EWTS syllable.
    """
    syllables = tokenize_ewts_syllables(
        text,
        strip_newlines=strip_newlines,
        strip_markers=strip_markers,
        strip_chars=strip_chars,
    )
    coord_map = [(start, end) for _, start, end in syllables]
    return [s for s, _, _ in syllables], coord_map


def encode_syllable_sequences(syllables_a, syllables_b, token_to_char):
    dense_a = "".join(token_to_char[s] for s in syllables_a)
    dense_b = "".join(token_to_char[s] for s in syllables_b)
    return dense_a, dense_b


def _syllable_k_to_dense_k(k: int, token_mode: str) -> int:
    """Convert syllable k-mer size to encoded dense-string k-mer size."""
    if token_mode == "syllable":
        return k * 2  # each syllable encodes to two PUA chars
    return k


def _syllable_units_to_dense(units: int, token_mode: str) -> int:
    """Convert syllable counts (gap/extend) to encoded dense-string units."""
    if token_mode == "syllable":
        return units * 2
    return units


def prepare_dense_sequences(text_a, text_b, token_mode="char", strip_chars=""):
    """
    Prepare dense alignment strings and coordinate maps for both texts.

    token_mode:
      - "char": align Latin characters (legacy)
      - "syllable": align EWTS syllable tokens
    """
    if token_mode == "syllable":
        syllables_a, map_a = build_syllable_alignment_map(text_a, strip_chars=strip_chars)
        syllables_b, map_b = build_syllable_alignment_map(text_b, strip_chars=strip_chars)
        token_to_char = build_syllable_encoding(syllables_a, syllables_b)
        dense_a, dense_b = encode_syllable_sequences(syllables_a, syllables_b, token_to_char)
        return dense_a, dense_b, map_a, map_b, "syllable", syllables_a, syllables_b

    dense_a, map_a = build_alignment_map(text_a, strip_chars=strip_chars)
    dense_b, map_b = build_alignment_map(text_b, strip_chars=strip_chars)
    return dense_a, dense_b, map_a, map_b, "char", None, None


def dense_range_to_original(coord_map, dense_start, dense_end, token_mode="char"):
    """
    Convert a [start, end) range in dense-string coordinates back to
    original-text coordinates.
    """
    if token_mode == "syllable":
        syl_start = dense_start // 2
        syl_end = max(syl_start + 1, (dense_end + 1) // 2)
        orig_start = coord_map[syl_start][0]
        orig_end = coord_map[syl_end - 1][1]
        return orig_start, orig_end

    orig_start = coord_map[dense_start]
    orig_end = coord_map[dense_end - 1] + 1
    return orig_start, orig_end


# ---------------------------------------------------------------------------
# 2. SEEDING — find exact k-mer matches
# ---------------------------------------------------------------------------

def build_kmer_index(text, k):
    """
    Build a dictionary mapping each k-mer to a list of start positions in `text`.
    """
    index = defaultdict(list)
    for i in range(len(text) - k + 1):
        kmer = text[i:i + k]
        index[kmer].append(i)
    return index


def find_seeds(dense_a, dense_b, k=15, max_kmer_hits=None):
    """
    Find all (pos_a, pos_b) pairs where dense_a[pos_a:pos_a+k] == dense_b[pos_b:pos_b+k].
    If max_kmer_hits is set, skip k-mers that appear more than that many times.
    """
    # Optimization: Build index on the shorter text to save RAM
    if len(dense_a) <= len(dense_b):
        short_text, long_text = dense_a, dense_b
        swapped = False
    else:
        short_text, long_text = dense_b, dense_a
        swapped = True

    index = build_kmer_index(short_text, k)

    # Pre-filter: remove k-mers that appear too many times in the short text
    if max_kmer_hits:
        index = {kmer: positions for kmer, positions in index.items()
                 if len(positions) <= max_kmer_hits}

    # Count k-mer frequencies in the long text to filter from that side too
    if max_kmer_hits:
        long_kmer_counts = defaultdict(int)
        for j in range(len(long_text) - k + 1):
            kmer = long_text[j:j + k]
            if kmer in index:
                long_kmer_counts[kmer] += 1
        # Remove k-mers that appear too many times in the long text
        for kmer, count in long_kmer_counts.items():
            if count > max_kmer_hits:
                del index[kmer]

    seeds = []
    for j in range(len(long_text) - k + 1):
        kmer = long_text[j:j + k]
        if kmer in index:
            for i in index[kmer]:
                if swapped:
                    seeds.append((j, i))
                else:
                    seeds.append((i, j))
    
    seeds.sort()
    return seeds


# ---------------------------------------------------------------------------
# 3. MERGING — The Critical Fix (Diagonal Lock + Slicer)
# ---------------------------------------------------------------------------

def merge_seeds_into_regions(seeds, k, max_gap=100, extend=200):
    """
    Group seeds into candidate regions, enforcing diagonal consistency to prevent 
    drifting into massive noise regions, and enforcing a size limit.
    """
    if not seeds:
        return []

    # 1. Sort primarily by Diagonal (Start - End), secondarily by Position
    # This naturally groups matches that lie on the same narrative "Lane".
    seeds.sort(key=lambda s: (s[0] - s[1], s[0]))

    clusters = []
    curr_seeds = [seeds[0]]
    ref_diag = seeds[0][0] - seeds[0][1]

    # --- SAFETY LIMIT: Cut region if it gets too long ---
    # 5,000 chars = ~25 million cells = ~1 second to process.
    MAX_REGION_LENGTH = 5000 

    for i in range(1, len(seeds)):
        seed = seeds[i]
        prev = curr_seeds[-1]
        
        curr_diag = seed[0] - seed[1]
        
        # 1. Gap Check (Connectivity)
        dist_ok = (seed[0] - prev[0]) <= max_gap
        
        # 2. Diagonal Check (Drift prevention)
        # Prevents "Daisy Chaining" across the text structure
        diag_ok = abs(curr_diag - ref_diag) <= max_gap
        
        # 3. Size Check (Hard Limit)
        # Prevents infinite regions even if the text matches perfectly
        len_ok = (seed[0] - curr_seeds[0][0]) < MAX_REGION_LENGTH

        if dist_ok and diag_ok and len_ok:
            curr_seeds.append(seed)
        else:
            clusters.append(curr_seeds)
            curr_seeds = [seed]
            ref_diag = curr_diag # Reset lane

    clusters.append(curr_seeds)

    regions = []
    for cluster in clusters:
        min_a = min(s[0] for s in cluster)
        max_a = max(s[0] for s in cluster) + k
        min_b = min(s[1] for s in cluster)
        max_b = max(s[1] for s in cluster) + k

        start_a = max(0, min_a - extend)
        end_a = max_a + extend
        start_b = max(0, min_b - extend)
        end_b = max_b + extend

        regions.append((start_a, end_a, start_b, end_b))

    return regions


# ---------------------------------------------------------------------------
# 4. ALIGNMENT — The Overflow Fix
# ---------------------------------------------------------------------------

def align_token_region(tokens_a, tokens_b, aligner, min_score=15.0, max_iterations=100):
    """Run iterative Smith-Waterman on syllable token lists."""
    work_a = list(tokens_a)
    work_b = list(tokens_b)
    MASK = object()
    matches = []

    for _ in range(max_iterations):
        active_a = [(j, t) for j, t in enumerate(work_a) if t is not MASK]
        active_b = [(j, t) for j, t in enumerate(work_b) if t is not MASK]
        if not active_a or not active_b:
            break

        vocab = sorted({t for _, t in active_a} | {t for _, t in active_b})
        enc = {t: chr(_PUA_START + i) for i, t in enumerate(vocab)}
        local_map_a = [j for j, _ in active_a]
        local_map_b = [j for j, _ in active_b]
        local_dense_a = "".join(enc[t] for _, t in active_a)
        local_dense_b = "".join(enc[t] for _, t in active_b)

        alignments = aligner.align(local_dense_a, local_dense_b)
        iterator = iter(alignments)
        try:
            best = next(iterator)
        except StopIteration:
            break

        if best.score < min_score:
            break

        ranges_a = best.aligned[0]
        ranges_b = best.aligned[1]
        ds_a, de_a = ranges_a[0][0], ranges_a[-1][1]
        ds_b, de_b = ranges_b[0][0], ranges_b[-1][1]

        work_start_a = local_map_a[ds_a]
        work_end_a = local_map_a[de_a - 1] + 1
        work_start_b = local_map_b[ds_b]
        work_end_b = local_map_b[de_b - 1] + 1

        matches.append((best.score, work_start_a, work_end_a, work_start_b, work_end_b))

        for j in range(work_start_a, work_end_a):
            work_a[j] = MASK
        for j in range(work_start_b, work_end_b):
            work_b[j] = MASK

    return matches


def align_region(dense_a, dense_b, aligner, min_score=15.0, max_iterations=100):
    """
    Run iterative Smith-Waterman on a pair of dense text regions.
    """
    work_a = list(dense_a)
    work_b = list(dense_b)
    MASK_CHAR = '\uFFF0'
    matches = []

    for _ in range(max_iterations):
        active_a = [(j, c) for j, c in enumerate(work_a) if c != MASK_CHAR]
        active_b = [(j, c) for j, c in enumerate(work_b) if c != MASK_CHAR]

        if not active_a or not active_b:
            break

        local_map_a = [j for j, c in active_a]
        local_dense_a = "".join(c for j, c in active_a)
        local_map_b = [j for j, c in active_b]
        local_dense_b = "".join(c for j, c in active_b)

        alignments = aligner.align(local_dense_a, local_dense_b)
        
        # --- FIX: Don't count alignments, just take the first ---
        # Prevents OverflowError on highly repetitive text
        iterator = iter(alignments)
        try:
            best = next(iterator)
        except StopIteration:
            break

        if best.score < min_score:
            break

        ranges_a = best.aligned[0]
        ranges_b = best.aligned[1]

        ds_a, de_a = ranges_a[0][0], ranges_a[-1][1]
        ds_b, de_b = ranges_b[0][0], ranges_b[-1][1]

        work_start_a = local_map_a[ds_a]
        work_end_a = local_map_a[de_a - 1] + 1
        work_start_b = local_map_b[ds_b]
        work_end_b = local_map_b[de_b - 1] + 1

        matches.append((best.score, work_start_a, work_end_a, work_start_b, work_end_b))

        # Mask matched characters
        for j in range(work_start_a, work_end_a):
            work_a[j] = MASK_CHAR
        for j in range(work_start_b, work_end_b):
            work_b[j] = MASK_CHAR

    return matches


# ---------------------------------------------------------------------------
# 5. MAIN PIPELINE
# ---------------------------------------------------------------------------

def seed_and_extend(text_a, text_b, match_score=1.0, mismatch_score=-1.5,
                    open_gap_score=-1.0, extend_gap_score=-1.0,
                    min_score=15.0, max_iterations=100,
                    seed_k=15, seed_max_gap=100, seed_extend=200,
                    seed_max_kmer_hits=None, strip_chars="", token_mode="syllable",
                    quiet=False):

    dense_a, dense_b, map_a, map_b, token_mode, syllables_a, syllables_b = prepare_dense_sequences(
        text_a, text_b, token_mode=token_mode, strip_chars=strip_chars,
    )
    unit = "syllables" if token_mode == "syllable" else "chars"

    if not quiet:
        print(f"Text A: {len(text_a)} original -> {len(dense_a)} dense {unit}")
        print(f"Text B: {len(text_b)} original -> {len(dense_b)} dense {unit}")
        print("Finding seeds...")
    dense_k = _syllable_k_to_dense_k(seed_k, token_mode)
    seeds = find_seeds(dense_a, dense_b, k=dense_k, max_kmer_hits=seed_max_kmer_hits)
    if not quiet:
        print(f"Found {len(seeds)} seed hits")

    if not seeds:
        return []

    regions = merge_seeds_into_regions(
        seeds,
        k=dense_k,
        max_gap=_syllable_units_to_dense(seed_max_gap, token_mode),
        extend=_syllable_units_to_dense(seed_extend, token_mode),
    )
    
    # Clamp to text boundaries
    regions = [
        (sa, min(ea, len(dense_a)), sb, min(eb, len(dense_b)))
        for sa, ea, sb, eb in regions
    ]
    if not quiet:
        print(f"Merged into {len(regions)} candidate regions")
        sizes = [(ea - sa) * (eb - sb) for sa, ea, sb, eb in regions]
        sizes.sort(reverse=True)
        if sizes:
            print(f"Top 5 region sizes (cells): {sizes[:5]}")

    aligner = PairwiseAligner()
    aligner.mode = 'local'
    aligner.match_score = match_score
    aligner.mismatch_score = mismatch_score
    aligner.open_gap_score = open_gap_score
    aligner.extend_gap_score = extend_gap_score

    all_matches = []

    for sa, ea, sb, eb in tqdm(regions, desc="Aligning regions", disable=quiet):
        # SAFETY VALVE: Skip massive regions (just in case)
        area = (ea - sa) * (eb - sb)
        if area > 50_000_000: # 50 Million cells limit (approx 2-3 sec)
             # Optional: Log warning if skipping
             continue

        region_a = dense_a[sa:ea]
        region_b = dense_b[sb:eb]

        if token_mode == "syllable":
            region_tokens_a = syllables_a[sa // 2: ea // 2]
            region_tokens_b = syllables_b[sb // 2: eb // 2]
            region_matches = align_token_region(
                region_tokens_a, region_tokens_b, aligner,
                min_score=min_score,
                max_iterations=max_iterations,
            )
            for score, rs_a, re_a, rs_b, re_b in region_matches:
                glob_start_a = (sa // 2 + rs_a) * 2
                glob_end_a = (sa // 2 + re_a) * 2
                glob_start_b = (sb // 2 + rs_b) * 2
                glob_end_b = (sb // 2 + re_b) * 2
                orig_start_a, orig_end_a = dense_range_to_original(
                    map_a, glob_start_a, glob_end_a, token_mode=token_mode,
                )
                orig_start_b, orig_end_b = dense_range_to_original(
                    map_b, glob_start_b, glob_end_b, token_mode=token_mode,
                )
                match_seg_a = text_a[orig_start_a:orig_end_a]
                match_seg_b = text_b[orig_start_b:orig_end_b]
                all_matches.append((score, match_seg_a, match_seg_b,
                                    orig_start_a, orig_end_a, orig_start_b, orig_end_b))
            continue

        region_matches = align_region(region_a, region_b, aligner,
                                      min_score=min_score,
                                      max_iterations=max_iterations)

        for score, rs_a, re_a, rs_b, re_b in region_matches:
            glob_start_a = sa + rs_a
            glob_end_a = sa + re_a
            glob_start_b = sb + rs_b
            glob_end_b = sb + re_b

            orig_start_a, orig_end_a = dense_range_to_original(
                map_a, glob_start_a, glob_end_a, token_mode=token_mode,
            )
            orig_start_b, orig_end_b = dense_range_to_original(
                map_b, glob_start_b, glob_end_b, token_mode=token_mode,
            )

            match_seg_a = text_a[orig_start_a:orig_end_a]
            match_seg_b = text_b[orig_start_b:orig_end_b]

            all_matches.append((score, match_seg_a, match_seg_b,
                                orig_start_a, orig_end_a, orig_start_b, orig_end_b))

    # Remove overlapping matches, keeping higher-scoring ones
    all_matches.sort(key=lambda x: -x[0])
    filtered = []
    for match in all_matches:
        _, _, _, start_a, end_a, start_b, end_b = match
        overlaps = False
        for kept in filtered:
            _, _, _, ks_a, ke_a, ks_b, ke_b = kept
            ov_a = start_a < ke_a and ks_a < end_a
            ov_b = start_b < ke_b and ks_b < end_b
            if ov_a and ov_b:
                overlaps = True
                break
        if not overlaps:
            filtered.append(match)
    all_matches = filtered

    if not quiet:
        print(f"Found {len(all_matches)} total matches")
    all_matches.sort(key=lambda x: x[3])  # sort by start_a
    return all_matches


# ---------------------------------------------------------------------------
# 6. ORIGINAL FULL-TEXT FALLBACK (kept for legacy)
# ---------------------------------------------------------------------------

def smith_waterman_waterfall(text_a, text_b, match_score=1.0, mismatch_score=-1.5,
                             open_gap_score=-1.0, extend_gap_score=-1.0,
                             min_score=15.0, max_iterations=100,
                             strip_chars="", token_mode="syllable", quiet=False):
    aligner = PairwiseAligner()
    aligner.mode = 'local'
    aligner.match_score = match_score
    aligner.mismatch_score = mismatch_score
    aligner.open_gap_score = open_gap_score
    aligner.extend_gap_score = extend_gap_score

    matches = []
    dense_a, dense_b, map_a, map_b, token_mode, _syllables_a, _syllables_b = prepare_dense_sequences(
        text_a, text_b, token_mode=token_mode, strip_chars=strip_chars,
    )
    unit = "syllables" if token_mode == "syllable" else "chars"

    if not quiet:
        print(f"Text A: {len(text_a)} original chars -> {len(dense_a)} alignment {unit}")
        print(f"Text B: {len(text_b)} original chars -> {len(dense_b)} alignment {unit}")

    work_a = list(dense_a)
    work_b = list(dense_b)
    MASK_CHAR = '\uFFF0'

    pbar = tqdm(total=max_iterations, desc="Aligning", leave=False, disable=quiet)
    for i in range(max_iterations):
        pbar.update(1)

        iter_map_a = [map_a[j] for j, c in enumerate(work_a) if c != MASK_CHAR]
        iter_dense_a = "".join(c for c in work_a if c != MASK_CHAR)
        iter_map_b = [map_b[j] for j, c in enumerate(work_b) if c != MASK_CHAR]
        iter_dense_b = "".join(c for c in work_b if c != MASK_CHAR)

        alignments = aligner.align(iter_dense_a, iter_dense_b)
        
        # --- FIX: Overflow protection ---
        iterator = iter(alignments)
        try:
            best = next(iterator)
        except StopIteration:
            break
            
        if best.score < min_score:
            break

        pbar.set_postfix(score=f"{best.score:.1f}", matches=len(matches) + 1)

        ranges_a = best.aligned[0]
        ranges_b = best.aligned[1]
        dense_start_a, dense_end_a = ranges_a[0][0], ranges_a[-1][1]
        dense_start_b, dense_end_b = ranges_b[0][0], ranges_b[-1][1]

        orig_start_a, orig_end_a = dense_range_to_original(
            iter_map_a, dense_start_a, dense_end_a, token_mode=token_mode,
        )
        orig_start_b, orig_end_b = dense_range_to_original(
            iter_map_b, dense_start_b, dense_end_b, token_mode=token_mode,
        )

        match_seg_a = text_a[orig_start_a:orig_end_a]
        match_seg_b = text_b[orig_start_b:orig_end_b]
        matches.append((best.score, match_seg_a, match_seg_b,
                        orig_start_a, orig_end_a, orig_start_b, orig_end_b))

        # Simple masking (not index-mapped) for legacy waterfall
        idx = 0
        for j, c in enumerate(work_a):
            if c == MASK_CHAR: continue
            if dense_start_a <= idx < dense_end_a: work_a[j] = MASK_CHAR
            idx += 1
        idx = 0
        for j, c in enumerate(work_b):
            if c == MASK_CHAR: continue
            if dense_start_b <= idx < dense_end_b: work_b[j] = MASK_CHAR
            idx += 1

    pbar.close()
    return matches


# ---------------------------------------------------------------------------
# 6b. PARALLEL PAIR COMPARISON
# ---------------------------------------------------------------------------

def _init_compare_worker(documents, worker_config):
    _WORKER_STATE["documents"] = documents
    _WORKER_STATE["config"] = worker_config


def _matches_to_rows(file_a, file_b, matches):
    rows = []
    for score, text_a, text_b, start_a, end_a, start_b, end_b in matches:
        is_stock, phrase_type = classify_match_row(text_a, text_b)
        rows.append({
            "file_a": file_a,
            "file_b": file_b,
            "score": score,
            "text_a": text_a,
            "text_b": text_b,
            "start_a": start_a,
            "end_a": end_a,
            "start_b": start_b,
            "end_b": end_b,
            "len_a": end_a - start_a,
            "len_b": end_b - start_b,
            "stock_phrase": is_stock,
            "stock_phrase_type": phrase_type or "",
        })
    return rows


def _compare_pair(pair):
    file_a, file_b = pair
    documents = _WORKER_STATE["documents"]
    cfg = _WORKER_STATE["config"]

    if cfg["mode"] == "seed":
        matches = seed_and_extend(
            documents[file_a],
            documents[file_b],
            match_score=cfg["match_score"],
            mismatch_score=cfg["mismatch_score"],
            open_gap_score=cfg["open_gap_score"],
            extend_gap_score=cfg["extend_gap_score"],
            min_score=cfg["min_score"],
            max_iterations=cfg["max_iterations"],
            seed_k=cfg["seed_k"],
            seed_max_gap=cfg["seed_max_gap"],
            seed_extend=cfg["seed_extend"],
            seed_max_kmer_hits=cfg["seed_max_kmer_hits"],
            strip_chars=cfg["strip_chars"],
            token_mode=cfg["token_mode"],
            quiet=True,
        )
    else:
        matches = smith_waterman_waterfall(
            documents[file_a],
            documents[file_b],
            match_score=cfg["match_score"],
            mismatch_score=cfg["mismatch_score"],
            open_gap_score=cfg["open_gap_score"],
            extend_gap_score=cfg["extend_gap_score"],
            min_score=cfg["min_score"],
            max_iterations=cfg["max_iterations"],
            strip_chars=cfg["strip_chars"],
            token_mode=cfg["token_mode"],
            quiet=True,
        )

    return _matches_to_rows(file_a, file_b, matches)


def _compare_pair_serial(file_a, file_b, documents, args, algo, seed_k, seed_max_gap,
                         seed_extend, seed_max_kmer_hits, strip_chars, token_mode):
    print(f"\n--- {file_a} vs {file_b} ---")
    if args.mode == "seed":
        matches = seed_and_extend(
            documents[file_a],
            documents[file_b],
            match_score=algo["match_score"],
            mismatch_score=algo["mismatch_score"],
            open_gap_score=algo["open_gap_score"],
            extend_gap_score=algo["extend_gap_score"],
            min_score=algo["min_score"],
            max_iterations=algo["max_iterations"],
            seed_k=seed_k,
            seed_max_gap=seed_max_gap,
            seed_extend=seed_extend,
            seed_max_kmer_hits=seed_max_kmer_hits,
            strip_chars=strip_chars,
            token_mode=token_mode,
        )
    else:
        matches = smith_waterman_waterfall(
            documents[file_a],
            documents[file_b],
            match_score=algo["match_score"],
            mismatch_score=algo["mismatch_score"],
            open_gap_score=algo["open_gap_score"],
            extend_gap_score=algo["extend_gap_score"],
            min_score=algo["min_score"],
            max_iterations=algo["max_iterations"],
            strip_chars=strip_chars,
            token_mode=token_mode,
        )
    return _matches_to_rows(file_a, file_b, matches)


# ---------------------------------------------------------------------------
# 7. CLI ENTRY POINT
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", default="config.yaml")
    parser.add_argument("--mode", default="seed")
    parser.add_argument(
        "--workers",
        type=int,
        default=None,
        help="Number of parallel processes for pair comparisons (default: 1, or parallel.workers in config)",
    )
    args = parser.parse_args()

    with open(args.config, "r") as f:
        config = yaml.safe_load(f)

    algo = config["algorithm"]
    output_path = config["output"]["path"]
    stock_phrase_detection = config.get("output", {}).get("stock_phrase_detection", True)

    seed_cfg = config.get("seeding", {})
    seed_k = seed_cfg.get("k", 15)
    seed_max_gap = seed_cfg.get("max_gap", 100)
    seed_extend = seed_cfg.get("extend", 200)
    seed_max_kmer_hits = seed_cfg.get("max_kmer_hits", None)
    strip_chars = config.get("preprocessing", {}).get("strip_chars", "")
    token_mode = config.get("preprocessing", {}).get("token_mode", "syllable")
    if token_mode not in ("char", "syllable"):
        raise ValueError("preprocessing.token_mode must be 'char' or 'syllable'")
    workers = args.workers
    if workers is None:
        workers = config.get("parallel", {}).get("workers", 1)
    workers = max(1, int(workers))

    input_cfg = config["input"]
    documents = {}

    if "dir" in input_cfg:
        # Original mode: all-pairs from a single directory
        input_dir = input_cfg["dir"]
        txt_files = sorted([f for f in os.listdir(input_dir) if f.endswith(".txt")])
        for fname in txt_files:
            with open(os.path.join(input_dir, fname), "r", encoding="utf-8") as f:
                documents[fname] = f.read()
        pairs = list(itertools.combinations(txt_files, 2))
        print(f"Single-directory mode: {len(txt_files)} files, {len(pairs)} pairs (before skipping duplicates)")
    elif "point_of_comparison" in input_cfg and "corpus" in input_cfg:
        # Cross-comparison mode: point_of_comparison × corpus (recursive)
        poc_dir = input_cfg["point_of_comparison"]
        corpus_dir = input_cfg["corpus"]
        poc_files = sorted([f for f in os.listdir(poc_dir) if f.endswith(".txt")])
        corpus_files = []
        for root, dirs, files in os.walk(corpus_dir):
            dirs.sort()
            for f in sorted(files):
                if f.endswith(".txt"):
                    corpus_files.append(os.path.relpath(os.path.join(root, f), corpus_dir))
        for fname in poc_files:
            key = f"poc/{fname}"
            with open(os.path.join(poc_dir, fname), "r", encoding="utf-8") as f:
                documents[key] = f.read()
        for rel_path in corpus_files:
            key = f"corpus/{rel_path}"
            with open(os.path.join(corpus_dir, rel_path), "r", encoding="utf-8") as f:
                documents[key] = f.read()
        pairs = list(itertools.product(
            [f"poc/{f}" for f in poc_files],
            [f"corpus/{p}" for p in corpus_files],
        ))
        print(f"Cross-comparison mode: {len(poc_files)} point_of_comparison × {len(corpus_files)} corpus = {len(pairs)} pairs (before skipping duplicates)")
    else:
        raise ValueError("Config 'input' must specify either 'dir' or both 'point_of_comparison' and 'corpus'")

    n_before = len(pairs)
    pairs = [
        (a, b)
        for a, b in pairs
        if documents[a] != documents[b]
    ]
    n_skipped = n_before - len(pairs)
    if n_skipped:
        print(f"Skipping {n_skipped} pair(s) where both documents are identical (same full text).")
    print(f"Comparing {len(pairs)} pair(s) with {workers} worker(s), token_mode={token_mode}.")

    all_matches = []
    worker_config = {
        "mode": args.mode,
        "match_score": algo["match_score"],
        "mismatch_score": algo["mismatch_score"],
        "open_gap_score": algo["open_gap_score"],
        "extend_gap_score": algo["extend_gap_score"],
        "min_score": algo["min_score"],
        "max_iterations": algo["max_iterations"],
        "seed_k": seed_k,
        "seed_max_gap": seed_max_gap,
        "seed_extend": seed_extend,
        "seed_max_kmer_hits": seed_max_kmer_hits,
        "strip_chars": strip_chars,
        "token_mode": token_mode,
    }

    if workers == 1:
        for file_a, file_b in tqdm(pairs, desc="Comparing pairs"):
            all_matches.extend(
                _compare_pair_serial(
                    file_a, file_b, documents, args, algo,
                    seed_k, seed_max_gap, seed_extend, seed_max_kmer_hits, strip_chars, token_mode,
                )
            )
    else:
        with ProcessPoolExecutor(
            max_workers=workers,
            initializer=_init_compare_worker,
            initargs=(documents, worker_config),
        ) as pool:
            for rows in tqdm(
                pool.map(_compare_pair, pairs, chunksize=1),
                total=len(pairs),
                desc="Comparing pairs",
            ):
                all_matches.extend(rows)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    if not stock_phrase_detection:
        for row in all_matches:
            row["stock_phrase"] = False
            row["stock_phrase_type"] = ""
    df = pd.DataFrame(all_matches, columns=["file_a", "file_b", "score",
                                            "text_a", "text_b",
                                            "start_a", "end_a",
                                            "start_b", "end_b",
                                            "len_a", "len_b",
                                            "stock_phrase", "stock_phrase_type"])
    df = df.sort_values(["file_a", "file_b", "start_a"], ignore_index=True)
    df.to_csv(output_path, index=False)
    print(f"\nWrote {len(df)} results to {output_path}")

    # Write a Word document with a review table
    docx_path = os.path.splitext(output_path)[0] + ".docx"
    _write_review_docx(df, docx_path)


def _write_review_docx(df, docx_path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)

    stock_fill = parse_xml(
        r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w'))
    )

    for (file_a, file_b), group in df.groupby(["file_a", "file_b"], sort=False):
        doc.add_heading(f"{file_a}  ↔  {file_b}", level=2)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(0.4)
        table.columns[1].width = Inches(0.7)
        table.columns[2].width = Inches(0.9)
        table.columns[3].width = Inches(2.5)
        table.columns[4].width = Inches(2.5)

        hdr = table.rows[0].cells
        for cell, text in zip(hdr, ["#", "Stock phrase", "Type", file_a, file_b]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True

        for idx, (_, row) in enumerate(group.iterrows(), 1):
            cells = table.add_row().cells
            is_stock = bool(row.get("stock_phrase", False))
            phrase_type = str(row.get("stock_phrase_type", "") or "")
            cells[0].text = str(idx)
            cells[1].text = "Yes" if is_stock else "No"
            cells[2].text = phrase_type
            cells[3].text = str(row["text_a"])
            cells[4].text = str(row["text_b"])
            if is_stock:
                for cell in cells:
                    cell._tc.get_or_add_tcPr().append(stock_fill)

        doc.add_paragraph("")

    doc.save(docx_path)
    print(f"Wrote review document to {docx_path}")


if __name__ == "__main__":
    main()